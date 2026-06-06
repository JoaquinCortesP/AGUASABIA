from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
INPUT_MD = ROOT / "Documentacion" / "docs_tecnicos" / "guia_postman_backend_aguasabia.md"
OUTPUT_DOCX = ROOT / "Documentacion" / "docs_tecnicos" / "guia_postman_backend_aguasabia.docx"


def set_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    for style_name, size, color in (
        ("Title", 22, RGBColor(14, 74, 114)),
        ("Heading 1", 16, RGBColor(19, 86, 133)),
        ("Heading 2", 13, RGBColor(34, 34, 34)),
        ("Heading 3", 11, RGBColor(52, 52, 52)),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A7C7DC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    run = title.add_run("Guia Backend AguaSabia para Postman")
    run.bold = True
    title.paragraph_format.space_after = Pt(4)
    add_paragraph_border(title)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Revision tecnica del backend, ajustes para Postman y hoja de ruta de integracion con Leaflet, Open-Meteo y tecnologias satelitales."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(84, 84, 84)


def add_code_block(document: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def build_docx(source_text: str) -> Document:
    document = Document()
    set_page_geometry(document)
    style_document(document)
    add_title_block(document)

    lines = source_text.splitlines()
    code_buffer: list[str] = []
    in_code_block = False

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("# Guia Backend AguaSabia para Postman"):
            continue

        if line.startswith("```"):
            if in_code_block:
                add_code_block(document, code_buffer)
                code_buffer = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue

        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
            continue

        if line.startswith("- "):
            add_bullet(document, line[2:])
            continue

        if len(line) > 3 and line[0].isdigit() and line[1] == "." and line[2] == " ":
            add_numbered(document, line[3:])
            continue

        add_body_paragraph(document, line)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 1
    footer_run = footer.add_run("AguaSabia - Guia tecnica backend")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(110, 110, 110)
    footer_run.add_break(WD_BREAK.LINE)
    footer_run.add_text("Actualizada el 27-05-2026")

    return document


def main() -> None:
    source_text = INPUT_MD.read_text(encoding="utf-8")
    document = build_docx(source_text)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Documento generado en: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
